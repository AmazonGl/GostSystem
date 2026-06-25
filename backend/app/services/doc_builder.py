import os
import sys
import re
import subprocess
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Глобальные настройки оформления основного текста (задаются при сборке)
_FMT = {
    "font": "Times New Roman",
    "size": 14,
    "align": "justify",   # justify | left | center | right
    "line_spacing": 1.5,
}

_ALIGN_MAP = {
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _body_align():
    return _ALIGN_MAP.get(_FMT["align"], WD_ALIGN_PARAGRAPH.JUSTIFY)
from app.config import settings


def _set_page_margins(doc):
    for section in doc.sections:
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def _set_font(run, size=None, bold=False, italic=False):
    run.font.name = _FMT["font"]
    run.font.size = Pt(size if size is not None else _FMT["size"])
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def _clean(text: str) -> str:
    """Убираем все markdown символы"""
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def _add_title_page(doc, doc_title: str, gost_code: str, org_name: str = "", title_page: dict | None = None):
    tp = title_page or {}
    year = str(datetime.now().year)

    # Организация-разработчик
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tp.get("org_name") or org_name or "ОРГАНИЗАЦИЯ-РАЗРАБОТЧИК")
    _set_font(r, 12)

    # Исполнитель (необязательно)
    executor = tp.get("executor", "")
    if executor:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(executor)
        _set_font(r, 12)

    for _ in range(3):
        doc.add_paragraph()

    # Гриф утверждения
    approve_label = tp.get("approve_label", "УТВЕРЖДАЮ")
    approve_position = tp.get("approve_position", "")
    approve_name = tp.get("approve_name", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(approve_label)
    _set_font(r, 12, bold=True)

    if approve_position:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(approve_position)
        _set_font(r, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(approve_name if approve_name else "________________________")
    _set_font(r, 12)

    approve_date = tp.get("approve_date", "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(approve_date if approve_date else f'«___» ____________ {year} г.')
    _set_font(r, 12)

    for _ in range(4):
        doc.add_paragraph()

    # Название документа
    title_text = tp.get("doc_title") or doc_title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_text.upper())
    _set_font(r, 18, bold=True)

    # Шифр / обозначение
    cipher = tp.get("cipher", "")
    if cipher:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cipher)
        _set_font(r, 14)

    gost_display = tp.get("gost_code", "") or gost_code
    if gost_display:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(gost_display)
        _set_font(r, 14)

    # Стадия / версия
    stage = tp.get("stage", "")
    if stage:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(stage)
        _set_font(r, 14)

    for _ in range(6):
        doc.add_paragraph()

    # Город и год
    city = tp.get("city", "")
    doc_year = tp.get("year", "") or year
    footer_text = f"{city + '  ' if city else ''}{doc_year}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(footer_text)
    _set_font(r, 14)

    doc.add_page_break()


def _add_toc(doc, structure: list):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("СОДЕРЖАНИЕ")
    _set_font(r, 14, bold=True)
    doc.add_paragraph()

    for i, item in enumerate(structure):
        title = _clean(item.get("section", f"Раздел {i+1}"))
        dots = max(1, 55 - len(title))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(f"{title} {'.' * dots} {i + 3}")
        _set_font(r, 14)

    doc.add_page_break()


def _add_section_heading(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(_clean(text))
    _set_font(r, 14, bold=True)


def _add_body(doc, text: str):
    """Парсим текст и добавляем правильно отформатированные абзацы"""
    lines = text.split('\n')
    n = len(lines)
    idx = 0
    while idx < n:
        line = lines[idx].strip()
        idx += 1
        if not line:
            continue

        # Картинка: [[IMG:путь|Название]]
        m_img = re.match(r'^\[\[IMG:(.+?)\|(.*?)\]\]$', line)
        if m_img:
            img_path, caption = m_img.group(1).strip(), m_img.group(2).strip()
            if os.path.exists(img_path):
                # счётчик рисунков храним на объекте документа
                num = getattr(doc, "_figure_counter", 0) + 1
                doc._figure_counter = num
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_img.add_run()
                try:
                    run.add_picture(img_path, width=Cm(15))
                except Exception:
                    pass
                # подпись "Рисунок N — Название"
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_text = f"Рисунок {num} — {caption}" if caption else f"Рисунок {num}"
                rc = p_cap.add_run(cap_text)
                _set_font(rc, 14)
                p_cap.paragraph_format.space_after = Pt(10)
            continue

        # Заголовок markdown ## ###
        if re.match(r'^#{1,6}\s+', line):
            clean = _clean(line)
            if not clean:
                continue
            # Пропускаем пустой подзаголовок — без текста до следующего заголовка
            has_body = False
            for look in range(idx, n):
                nxt = lines[look].strip()
                if not nxt:
                    continue
                if re.match(r'^#{1,6}\s+', nxt):
                    break
                has_body = True
                break
            if not has_body:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(clean)
            _set_font(r, 14, bold=True)

        # Маркированный список * - •
        elif re.match(r'^[\*\-•]\s+', line):
            clean = _clean(re.sub(r'^[\*\-•]\s+', '', line))
            if not clean:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing = _FMT["line_spacing"]
            r = p.add_run(f"– {clean}")
            _set_font(r)

        # Нумерованный список 1. 2. 3)
        elif re.match(r'^\d+[\.\)]\s+', line):
            clean = _clean(line)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing = _FMT["line_spacing"]
            r = p.add_run(clean)
            _set_font(r)

        # Строка только из ** (подзаголовок)
        elif re.match(r'^\*\*.+\*\*:?$', line):
            clean = _clean(line)
            if not clean:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(clean)
            _set_font(r, 14, bold=True)

        # Обычный текст
        else:
            clean = _clean(line)
            if not clean:
                continue
            p = doc.add_paragraph()
            p.alignment = _body_align()
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.line_spacing = _FMT["line_spacing"]
            r = p.add_run(clean)
            _set_font(r)


def build_docx(
    roadmap_structure: list,
    answers: list,
    generated_texts: list,
    doc_title: str = "Технический документ",
    gost_code: str = "",
    org_name: str = "",
    fmt: dict | None = None,
    title_page: dict | None = None,
) -> str:
    # Применяем пользовательские настройки оформления
    if fmt:
        _FMT["font"] = fmt.get("font") or _FMT["font"]
        _FMT["size"] = fmt.get("size") or _FMT["size"]
        _FMT["align"] = fmt.get("align") or _FMT["align"]
        _FMT["line_spacing"] = fmt.get("line_spacing") or _FMT["line_spacing"]
    else:
        # сброс к значениям по умолчанию
        _FMT.update({"font": "Times New Roman", "size": 14, "align": "justify", "line_spacing": 1.5})

    doc = Document()
    _set_page_margins(doc)

    style = doc.styles["Normal"]
    style.font.name = _FMT["font"]
    style.font.size = Pt(_FMT["size"])

    _add_title_page(doc, doc_title, gost_code, org_name, title_page)
    _add_toc(doc, roadmap_structure)

    for i, item in enumerate(roadmap_structure):
        section_title = item.get("section", f"Раздел {i + 1}")
        # Каждый новый раздел начинается с новой страницы (кроме первого)
        if i > 0:
            doc.add_page_break()
        _add_section_heading(doc, section_title)

        text = generated_texts[i] if i < len(generated_texts) else answers[i] if i < len(answers) else ""
        if text:
            _add_body(doc, text)

    out_dir = os.path.join(settings.STORAGE_PATH, "_generated")
    os.makedirs(out_dir, exist_ok=True)
    filename = f"{uuid.uuid4()}.docx"
    path = os.path.join(out_dir, filename)
    doc.save(path)
    return path


def convert_to_pdf(docx_path: str) -> str:
    """Конвертирует .docx в PDF через LibreOffice, обновляя оглавление (TOC)
    и поля документа — иначе содержание в PDF будет пустым.

    Способ: поднимаем LibreOffice как UNO-сервер (socket), открываем документ,
    обновляем индексы/поля программно и экспортируем PDF. Это надёжнее запуска
    Basic-макроса из командной строки (тот часто молча не выполняется).
    Если UNO по какой-то причине недоступен — откатываемся на обычную конвертацию.
    """
    import tempfile, shutil
    out_dir = os.path.dirname(docx_path)
    pdf_path = docx_path.replace(".docx", ".pdf")
    binary = shutil.which("soffice") or shutil.which("libreoffice") or "libreoffice"

    # 1) Надёжный путь — обновление TOC через UNO
    try:
        if _convert_with_uno(docx_path, pdf_path, binary):
            return pdf_path
    except Exception:
        pass

    # 2) Фолбэк: обычная конвертация (TOC может остаться пустым, но PDF будет)
    profile = tempfile.mkdtemp(prefix="lo_profile_")
    try:
        result = subprocess.run(
            [binary, f"-env:UserInstallation=file://{profile}",
             "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            capture_output=True, text=True, timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice error: {result.stderr or result.stdout}")
        if not os.path.exists(pdf_path):
            raise RuntimeError("PDF файл не был создан")
        return pdf_path
    finally:
        shutil.rmtree(profile, ignore_errors=True)


def update_docx_toc(docx_path: str) -> bool:
    """Обновляет оглавление (TOC) прямо в .docx через UNO и сохраняет файл обратно,
    чтобы содержание было готовым сразу в Word (без нажатия F9). True при успехе."""
    import shutil, logging
    log = logging.getLogger("toc")
    binary = shutil.which("soffice") or shutil.which("libreoffice") or "libreoffice"
    try:
        ok = _uno_process(docx_path, docx_path=docx_path, pdf_path=None, binary=binary)
        if not ok:
            log.warning("update_docx_toc: НЕ удалось обновить TOC через UNO для %s "
                        "(оглавление останется с пометкой F9). Проверьте, что в контейнере "
                        "установлен python3-uno и доступен системный python с модулем uno.",
                        docx_path)
        else:
            log.info("update_docx_toc: оглавление обновлено для %s", docx_path)
        return ok
    except Exception as e:
        log.warning("update_docx_toc: исключение %s", e)
        return False


def _convert_with_uno(docx_path: str, pdf_path: str, binary: str) -> bool:
    """Открывает документ через UNO, обновляет оглавление и поля, экспортирует PDF."""
    return _uno_process(docx_path, docx_path=None, pdf_path=pdf_path, binary=binary)


def _uno_process(src_path: str, docx_path: str | None, pdf_path: str | None, binary: str) -> bool:
    """Обновляет оглавление/поля документа и сохраняет в DOCX и/или PDF.

    Запускает LibreOffice как UNO-сокет-сервер и подключается к нему. Для импорта
    `uno` используется питон, у которого есть доступ к UNO: сперва пробуем текущий
    интерпретатор (работает, если установлен python3-uno совместимой версии), а в
    качестве запасного варианта — питон, поставляемый с LibreOffice.
    """
    import tempfile, shutil, time, socket, json

    profile = tempfile.mkdtemp(prefix="lo_uno_")

    def _free_port() -> int:
        s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        s.bind(("localhost", 0)); port = s.getsockname()[1]; s.close()
        return port

    port = _free_port()
    server = subprocess.Popen(
        [binary, f"-env:UserInstallation=file://{profile}",
         "--headless", "--invisible", "--nodefault", "--norestore", "--nologo",
         f"--accept=socket,host=localhost,port={port};urp;"],
        stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
    )

    # воркер-скрипт подключается к серверу и делает работу
    worker = os.path.join(profile, "_toc_worker.py")
    with open(worker, "w", encoding="utf-8") as f:
        f.write(_TOC_WORKER_SCRIPT)

    def _url(p):
        return "file://" + os.path.abspath(p)

    worker_args = [
        worker, str(port), _url(src_path),
        _url(docx_path) if docx_path else "-",
        _url(pdf_path) if pdf_path else "-",
    ]

    # питоны-кандидаты для воркера: нужен интерпретатор с доступом к `uno`.
    # В Docker (python:slim) питон приложения (sys.executable, /usr/local/bin)
    # обычно НЕ имеет uno, а Debian-пакет python3-uno ставит модуль под
    # СИСТЕМНЫЙ python (/usr/bin/python3.X). Поэтому пробуем явные системные пути.
    candidates = []
    lo_python = _find_libreoffice_python(binary)
    if lo_python:
        candidates.append(lo_python)
    candidates += [
        "/usr/bin/python3",
        "/usr/bin/python3.13", "/usr/bin/python3.12", "/usr/bin/python3.11",
        "/usr/bin/python3.10", "/usr/bin/python3.9",
        sys.executable, "python3",
    ]
    # убираем дубли, сохраняя порядок
    seen = set(); candidates = [c for c in candidates if c and not (c in seen or seen.add(c))]

    try:
        for py in candidates:
            if not py:
                continue
            try:
                r = subprocess.run([py, *worker_args], capture_output=True, text=True, timeout=180)
            except Exception:
                continue
            ok = True
            if docx_path:
                ok = ok and os.path.exists(docx_path)
            if pdf_path:
                ok = ok and os.path.exists(pdf_path)
            # воркер печатает MARK при успешном обновлении TOC
            if "TOC_UPDATED_OK" in (r.stdout or "") and ok:
                return True
        return False
    finally:
        try:
            server.terminate(); server.wait(timeout=10)
        except Exception:
            server.kill()
        shutil.rmtree(profile, ignore_errors=True)


def _find_libreoffice_python(binary: str):
    """Ищет интерпретатор Python, поставляемый с LibreOffice (имеет доступ к uno)."""
    import shutil as _sh
    for cand in ("/usr/lib/libreoffice/program/python",
                 "/opt/libreoffice/program/python",
                 "/usr/bin/libreoffice-python"):
        if os.path.exists(cand):
            return cand
    return None


# Воркер: подключается к запущенному UNO-серверу, обновляет TOC/поля и сохраняет.
# Запускается интерпретатором, у которого есть `uno` (питон LibreOffice или python3-uno).
_TOC_WORKER_SCRIPT = r'''
import sys, time
try:
    import uno
    from com.sun.star.beans import PropertyValue
except Exception:
    sys.exit(0)

def prop(name, value):
    p = PropertyValue(); p.Name = name; p.Value = value
    return p

def run():
    port, src_url, docx_url, pdf_url = sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4]
    local_ctx = uno.getComponentContext()
    resolver = local_ctx.ServiceManager.createInstanceWithContext(
        "com.sun.star.bridge.UnoUrlResolver", local_ctx)
    ctx = None
    for _ in range(60):
        try:
            ctx = resolver.resolve(
                "uno:socket,host=localhost,port=%s;urp;StarOffice.ComponentContext" % port)
            break
        except Exception:
            time.sleep(0.5)
    if ctx is None:
        return
    smgr = ctx.ServiceManager
    desktop = smgr.createInstanceWithContext("com.sun.star.frame.Desktop", ctx)
    doc = desktop.loadComponentFromURL(src_url, "_blank", 0, (prop("Hidden", True),))
    if doc is None:
        return
    try:
        idx = doc.getDocumentIndexes()
        for i in range(idx.getCount()):
            idx.getByIndex(i).update()
    except Exception:
        pass
    try: doc.getTextFields().refresh()
    except Exception: pass
    try: doc.refresh()
    except Exception: pass
    if docx_url and docx_url != "-":
        doc.storeToURL(docx_url, (prop("FilterName", "MS Word 2007 XML"),))
    if pdf_url and pdf_url != "-":
        doc.storeToURL(pdf_url, (prop("FilterName", "writer_pdf_Export"),))
    doc.close(False)
    print("TOC_UPDATED_OK")

try:
    run()
except Exception:
    pass
'''


