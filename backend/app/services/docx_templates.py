"""Работа с загружаемыми .docx-шаблонами документов.

Подход «по заголовкам»: шаблон — обычный .docx с оформленными заголовками
разделов (стили Heading) и текстом-рыбой под ними. Система:
  1. parse_template_structure — читает заголовки и строит дерево разделов;
  2. render_template — вставляет текст пользователя под нужные заголовки,
     удаляя прежнее содержимое раздела (рыбу), и заменяет {{метки}} на титуле.
Оформление шаблона (шрифты, отступы, титул, стили) полностью сохраняется.
"""
import os
import re
import copy
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _is_heading(par) -> int:
    """Возвращает уровень заголовка (1..9) или 0, если абзац не заголовок.

    Распознаёт три случая:
      1) стиль Heading N / «Заголовок N»;
      2) ручная нумерация в начале строки (1, 1.1, 3.5.1);
      3) короткая строка целиком жирным начертанием.
    """
    style = (par.style.name or "") if par.style else ""
    m = re.match(r"Heading\s*(\d+)", style)
    if m:
        return int(m.group(1))
    m = re.match(r"Заголовок\s*(\d+)", style)
    if m:
        return int(m.group(1))

    text = (par.text or "").strip()
    if not text or len(text) > 90:
        return 0

    # Ручная нумерация: уровень = число сегментов (1 -> 1, 1.2 -> 2, 1.2.3 -> 3)
    m = re.match(r"^(\d+(?:\.\d+){0,3})\.?\s+\S", text)
    if m:
        depth = m.group(1).count(".") + 1
        return min(depth, 9)

    # Полностью жирная короткая строка без завершающей точки — вероятный заголовок
    runs = [r for r in par.runs if (r.text or "").strip()]
    if runs and all(r.bold for r in runs) and not text.endswith(".") and len(text.split()) <= 8:
        return 1

    return 0


def parse_template_structure(file_path: str) -> dict:
    """Строит вложенную структуру разделов шаблона из заголовков любой глубины.

    Возвращает {"sections": [{id, title, level, subsections: [...]}, ...]},
    где subsections могут содержать свои subsections (до 4+ уровней).
    Номер (id) строится по позиции в иерархии: 5, 5.5, 5.5.2, 5.5.2.1 и т.д.
    """
    doc = Document(file_path)
    flat = []
    for par in doc.paragraphs:
        lvl = _is_heading(par)
        title = par.text.strip()
        if lvl and title:
            flat.append({"title": title, "level": lvl})

    root = {"subsections": [], "level": 0}
    # стек узлов по уровням: stack[k] — последний открытый узел уровня k
    stack = [root]
    counters = {}  # уровень -> текущий счётчик на этом уровне в рамках родителя

    for item in flat:
        lvl = item["level"]
        # поднимаемся по стеку до родителя нужного уровня
        while len(stack) > lvl:
            stack.pop()
        # если перепрыгнули уровень (например с 1 сразу на 3) — нормализуем
        if len(stack) < lvl:
            lvl = len(stack)
        parent = stack[-1]

        # номер узла = номер родителя + порядковый индекс среди его детей
        index = len(parent["subsections"]) + 1
        parent_id = parent.get("id")
        node_id = f"{parent_id}.{index}" if parent_id else str(index)

        node = {"id": node_id, "title": item["title"], "level": lvl, "subsections": []}
        parent["subsections"].append(node)
        # этот узел становится текущим на своём уровне
        if len(stack) == lvl:
            stack.append(node)
        else:
            stack[lvl] = node
            del stack[lvl + 1:]

    return {"sections": root["subsections"]}


def _clear_section_body(doc, start_idx: int, end_idx: int):
    """Удаляет абзацы тела раздела (между start_idx+1 и end_idx, не трогая заголовки)."""
    body = doc.paragraphs
    to_remove = []
    for i in range(start_idx + 1, end_idx):
        # не удаляем вложенные заголовки
        if _is_heading(body[i]) == 0:
            to_remove.append(body[i])
    for p in to_remove:
        p._element.getparent().remove(p._element)


def _insert_paragraph_after(paragraph, text="", style=None):
    """Вставляет новый ЧИСТЫЙ абзац после указанного, возвращает его.

    Важно: не копируем исходный абзац (часто это заголовок) целиком, иначе новый
    абзац унаследует прямое форматирование заголовка (выравнивание, отступы, жирность)
    и перебьёт стиль Normal шаблона. Создаём пустой <w:p> — он возьмёт оформление
    из стиля, который зададим ниже.
    """
    from docx.oxml.ns import qn as _qn
    new_p = paragraph._p.makeelement(_qn('w:p'), {})
    paragraph._p.addnext(new_p)
    np = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            np.style = style
        except Exception:
            pass
    if text:
        np.add_run(text)
    return np


def _format_body_paragraph(par, is_list=False):
    """Оформление абзаца тела наследуется от стиля Normal шаблона — мы НЕ навязываем
    свои отступы/выравнивание, чтобы текст выглядел ровно как задано в шаблоне
    (обычно выравнивание по ширине + красная строка из стиля Normal).

    Делаем только минимум:
    - стиль Normal (чтобы не наследовать оформление заголовка);
    - для пунктов перечня — маркер-тире в начале (как принято в ЕСПД).
    Прямые свойства абзаца НЕ задаём — пусть берутся из стиля шаблона.
    """
    try:
        par.style = "Normal"
    except Exception:
        pass
    if is_list:
        if par.runs and not par.runs[0].text.lstrip().startswith(("–", "—", "-", "•")):
            par.runs[0].text = "– " + par.runs[0].text.lstrip()
    return par


def _insert_image_after(paragraph, img_path: str, caption: str, fig_num: int):
    """Вставляет картинку (по центру) с подписью «Рисунок N — …» после абзаца.

    Перед картинкой добавляется пустая строка для отступа от текста.
    Возвращает последний вставленный абзац (подпись), чтобы продолжать вставку.
    """
    # пустая строка-отступ перед рисунком
    spacer = _insert_paragraph_after(paragraph)
    # абзац с картинкой
    img_par = _insert_paragraph_after(spacer)
    img_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_par.add_run()
    try:
        run.add_picture(img_path, width=Cm(15))
    except Exception:
        # если файл не читается — оставляем пустой абзац, не падаем
        return img_par
    # подпись
    cap_par = _insert_paragraph_after(img_par)
    cap_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_text = f"Рисунок {fig_num} — {caption}" if caption else f"Рисунок {fig_num}"
    cap_par.add_run(cap_text)
    return cap_par


def _refresh_toc_fields(doc):
    """Очищает закэшированные строки оглавления (TOC) и помечает поля документа
    на пересчёт, чтобы Word/LibreOffice пересобрали оглавление при открытии.
    Без этого в готовом документе остаётся старый список разделов из шаблона.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 1) updateFields=true в settings — обновить поля при открытии
    try:
        settings = doc.part.document.settings.element
        if settings.find(qn("w:updateFields")) is None:
            upd = OxmlElement("w:updateFields")
            upd.set(qn("w:val"), "true")
            settings.append(upd)
    except Exception:
        pass

    body = doc.element.body

    # 2) Находим sdt-контейнеры оглавления (в них инструкция TOC) и очищаем кэш.
    for sdt in body.findall(qn("w:sdt")):
        instrs = sdt.findall(".//" + qn("w:instrText"))
        is_toc = any("TOC" in (it.text or "") for it in instrs)
        if not is_toc:
            continue
        content = sdt.find(qn("w:sdtContent"))
        if content is None:
            continue
        paras = content.findall(qn("w:p"))
        if not paras:
            continue
        # Оставляем только первый абзац, в нём пересобираем чистое TOC-поле,
        # остальные (закэшированные строки оглавления) удаляем.
        keep = paras[0]
        for p in paras[1:]:
            content.remove(p)
        # вычищаем все run-ы первого абзаца
        for r in keep.findall(qn("w:r")):
            keep.remove(r)
        # строим: begin → instrText(TOC) → separate → "Обновите оглавление (F9)" → end
        def _mk_run():
            return OxmlElement("w:r")

        r_begin = _mk_run()
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
        fc.set(qn("w:dirty"), "true")  # MS Word пересоберёт оглавление при открытии
        r_begin.append(fc)

        r_instr = _mk_run()
        it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
        it.text = ' TOC \\o "1-3" \\h \\z \\u '
        r_instr.append(it)

        r_sep = _mk_run()
        sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
        r_sep.append(sep)

        r_txt = _mk_run()
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve")
        t.text = "Оглавление будет обновлено автоматически (F9)"
        r_txt.append(t)

        r_end = _mk_run()
        ec = OxmlElement("w:fldChar"); ec.set(qn("w:fldCharType"), "end")
        r_end.append(ec)

        for r in (r_begin, r_instr, r_sep, r_txt, r_end):
            keep.append(r)


def _is_in_heading_chain(par) -> bool:
    return _is_heading(par) > 0


def _norm(s: str) -> str:
    return re.sub(r"[^а-яёa-z0-9]", "", (s or "").lower())


def split_combined_section(section_title: str, text: str) -> dict:
    """Разбирает текст раздела со склеенными подразделами "## N.M Заголовок\\nтекст"
    обратно в {title, text, subsections}. Если маркеров "##" нет — текст идёт в тело раздела.
    """
    lines = (text or "").split("\n")
    head_re = re.compile(r"^\s*#{1,6}\s*(?:[\d.]+\s+)?(.+?)\s*$")
    subsections = []
    intro_lines = []
    cur = None
    for line in lines:
        if re.match(r"^\s*#{1,6}\s+", line):
            title = head_re.match(line).group(1).strip()
            cur = {"title": title, "text": ""}
            subsections.append(cur)
        elif cur is not None:
            cur["text"] += (("\n" if cur["text"] else "") + line)
        else:
            intro_lines.append(line)
    for s in subsections:
        s["text"] = s["text"].strip()
    return {"title": section_title, "text": "\n".join(intro_lines).strip(), "subsections": subsections}





def _replace_template_title_page(doc, tp: dict):
    """Заменяет титульный лист шаблона на сгенерированный из данных формы.

    Удаляет все абзацы до первого заголовка раздела (это титул шаблона) и
    вставляет новый титул из полей: организация, гриф утверждения, название
    документа, шифр, город, год. Учебные данные шаблона при этом убираются.
    """
    from datetime import datetime
    year = tp.get("year") or str(datetime.now().year)

    # Находим первый заголовок-раздел (Heading) — всё до него считаем титулом шаблона
    first_heading_idx = None
    for idx, par in enumerate(doc.paragraphs):
        if _is_heading(par):
            first_heading_idx = idx
            break
    if first_heading_idx is None:
        return  # нет заголовков — не трогаем

    first_heading = doc.paragraphs[first_heading_idx]

    # Удаляем ВСЁ содержимое титула шаблона до первого заголовка — и абзацы, и
    # ТАБЛИЦЫ (в учебных шаблонах шапка «СОГЛАСОВАНО/УТВЕРЖДАЮ» сделана таблицами).
    # Не трогаем блок оглавления (TOC), если он расположен до первого заголовка.
    from docx.oxml.ns import qn as _qn
    from lxml import etree
    fh_el = first_heading._p
    body = fh_el.getparent()
    toc_anchor = None  # первый сохранённый элемент TOC (перед ним вставим титул)
    for child in list(body):
        if child is fh_el:
            break  # дошли до первого заголовка — дальше тело документа
        if child.tag == _qn('w:sectPr'):
            continue
        # TOC / содержание: сохраняем и запоминаем как якорь
        try:
            xml = etree.tostring(child, encoding="unicode")
        except Exception:
            xml = ""
        is_toc = ("TOC" in xml) or ("w:sdt" in xml and "instrText" in xml)
        is_heading_word = child.tag == _qn('w:p') and ("Содержание" in xml or "Оглавление" in xml)
        if is_toc or is_heading_word:
            if toc_anchor is None:
                toc_anchor = child
            continue
        body.remove(child)

    # Точка вставки титула: перед оглавлением (если есть), иначе перед первым заголовком.
    anchor_el = toc_anchor if toc_anchor is not None else fh_el
    anchor_par = Paragraph(anchor_el, first_heading._parent)


    # Строим новый титул, вставляя абзацы ПЕРЕД первым заголовком
    def add_before(text, *, align=WD_ALIGN_PARAGRAPH.CENTER, bold=False, size=14, blank_after=0):
        p_el = copy.deepcopy(fh_el)
        for ch in list(p_el):
            p_el.remove(ch)
        anchor_el.addprevious(p_el)
        p = Paragraph(p_el, anchor_par._parent)
        try:
            p.style = "Normal"
        except Exception:
            pass
        p.alignment = align
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = None
        p.paragraph_format.line_spacing = 1.0
        if text:
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(size)
            r.font.bold = bold
        for _ in range(blank_after):
            b_el = copy.deepcopy(fh_el)
            for ch in list(b_el):
                b_el.remove(ch)
            anchor_el.addprevious(b_el)
        return p

    # Организация-разработчик (сверху по центру)
    add_before(tp.get("org_name") or "ОРГАНИЗАЦИЯ-РАЗРАБОТЧИК", bold=True, size=14)
    if tp.get("executor"):
        add_before(tp.get("executor"), size=12)
    add_before("", blank_after=2)

    # Гриф утверждения (справа)
    add_before(tp.get("approve_label") or "УТВЕРЖДАЮ", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size=12)
    if tp.get("approve_position"):
        add_before(tp.get("approve_position"), align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
    add_before(tp.get("approve_name") or "________________________", align=WD_ALIGN_PARAGRAPH.RIGHT, size=12)
    add_before(tp.get("approve_date") or f"«___» ____________ {year} г.", align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, blank_after=3)

    # Название документа (крупно по центру)
    add_before(tp.get("doc_title") or "ТЕХНИЧЕСКОЕ ЗАДАНИЕ", bold=True, size=20)
    if tp.get("cipher"):
        add_before(tp.get("cipher"), size=14, blank_after=1)
    if tp.get("gost_code"):
        add_before(tp.get("gost_code"), size=12)
    if tp.get("stage"):
        add_before(tp.get("stage"), size=12)
    add_before("", blank_after=4)

    # Город и год (внизу по центру)
    city_year = " ".join(x for x in [tp.get("city", ""), year] if x)
    add_before(city_year, size=12)

    # Явный разрыв страницы после титульного листа — титул занимает ровно 1 страницу,
    # дальше идёт содержание / первый раздел со следующей страницы.
    brk_p = add_before("")
    r_el = brk_p._p.makeelement(qn('w:r'), {})
    br_el = brk_p._p.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    r_el.append(br_el)
    brk_p._p.append(r_el)


def render_template(template_path: str, out_path: str, sections_data: list, title_vars: dict | None = None) -> str:
    """Вставляет текст пользователя в шаблон по заголовкам и сохраняет результат.

    sections_data: [{"title": "Введение", "text": "...", "subsections": [{"title","text"}]}]
    title_vars: значения для замены {{меток}} на титуле (название, год, ФИО и т.п.)
    """
    doc = Document(template_path)

    # 1) Титульный лист: если переданы данные титульника — генерируем свой
    #    (заменяя учебный титул шаблона), иначе заменяем {{метки}} если они есть.
    tp = title_vars or {}
    has_title_data = any(tp.get(k) for k in ("org_name", "doc_title", "approve_position", "approve_name", "cipher", "city", "executor"))
    if has_title_data:
        _replace_template_title_page(doc, tp)
    if title_vars:
        _replace_placeholders(doc, title_vars)

    # 2) Текст пользователя по заголовкам (плоско, ключ — нормализованное название)
    user_text = {}
    for sec in sections_data:
        if sec.get("text", "").strip():
            user_text[_norm(sec["title"])] = sec["text"]
        for sub in sec.get("subsections", []) or []:
            if sub.get("text", "").strip():
                user_text[_norm(sub["title"])] = sub["text"]

    # 3) Для каждого заголовка с текстом: удаляем старое тело (рыбу) между этим
    #    заголовком и следующим заголовком ЛЮБОГО уровня, затем вставляем текст.
    #    Работаем по ссылкам на абзацы, а не по индексам — безопасно при вставках.
    all_paras = list(doc.paragraphs)
    heads = [(p, _is_heading(p)) for p in all_paras]
    head_idxs = [k for k, (p, lvl) in enumerate(heads) if lvl]

    fig_counter = 0
    for pos, k in enumerate(head_idxs):
        par, lvl = heads[k]
        key = _norm(par.text.strip())
        if key not in user_text:
            continue
        # конец тела — следующий заголовок ЛЮБОГО уровня (тело = только прямой текст)
        next_head_k = head_idxs[pos + 1] if pos + 1 < len(head_idxs) else len(all_paras)
        # удаляем абзацы тела (между заголовком и следующим заголовком)
        for m in range(k + 1, next_head_k):
            victim = all_paras[m]
            if _is_heading(victim) == 0:
                victim._element.getparent().remove(victim._element)
        # вставляем пользовательский текст сразу после заголовка
        anchor = par
        for block in _split_blocks(user_text[key]):
            if block["type"] == "image":
                fig_counter += 1
                anchor = _insert_image_after(anchor, block["path"], block["caption"], fig_counter)
            else:
                is_list = block.get("style") == "List Bullet"
                anchor = _insert_paragraph_after(anchor, block["text"])
                _format_body_paragraph(anchor, is_list=is_list)

    # Заголовки оставляем как в шаблоне — их оформление (шрифт, размер, отбивка,
    # нумерация) задаётся стилями Heading самого шаблона.

    # Обновляем оглавление: очищаем устаревший кэш и помечаем поля на пересчёт
    _refresh_toc_fields(doc)

    # Очищаем метаданные шаблона (автор, компания и т.п. — например «Транснефть»),
    # чтобы они не попадали в итоговый документ и PDF.
    try:
        cp = doc.core_properties
        cp.author = ""
        cp.last_modified_by = ""
        cp.company = ""
        cp.title = title_vars.get("doc_title", "") if title_vars else ""
        cp.comments = ""
        cp.category = ""
        cp.keywords = ""
    except Exception:
        pass

    doc.save(out_path)
    return out_path


def _split_blocks(text: str):
    """Разбивает текст пользователя на абзацы/списки/картинки с учётом «кривого» ввода.

    Правила (предсказуемые, без угадывания):
    - пустая строка — граница абзаца;
    - строка с маркером (— - • * или «1.», «1)») — отдельный пункт списка;
    - подряд идущие обычные строки склеиваются в ОДИН абзац (мягкие переносы),
      чтобы разорванное на строки предложение не превращалось в кучу абзацев.
    """
    lines = text.split("\n")
    blocks = []
    para_buf = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            joined = " ".join(x.strip() for x in para_buf if x.strip())
            joined = re.sub(r'\s{2,}', ' ', joined).strip()  # схлопываем двойные пробелы
            if joined:
                blocks.append({"type": "text", "text": joined, "style": None})
            para_buf = []

    def is_marker(s):
        # маркер — тире/буллет (с пробелом или без: «– тест» и «-тест»), либо нумерация «1.»/«1)»
        return bool(re.match(r'^([—–\-•*]\s*\S|\d+[.)]\s+)', s))

    for raw in lines:
        s = raw.strip()
        if not s:
            flush_para()
            continue

        m_img = re.match(r'^\[\[IMG:(.+?)\|(.*?)\]\]$', s)
        if m_img:
            flush_para()
            blocks.append({"type": "image", "path": m_img.group(1).strip(), "caption": m_img.group(2).strip()})
            continue

        if is_marker(s):
            flush_para()
            item = re.sub(r'^([—–\-•*]\s*|\d+[.)]\s+)', '', s).strip()
            blocks.append({"type": "text", "text": item, "style": "List Bullet"})
            continue

        para_buf.append(s)

    flush_para()
    return blocks


def _replace_placeholders(doc, variables: dict):
    """Заменяет {{key}} и «Название программы»-подобные метки в тексте абзацев."""
    def repl_in_par(par):
        full = "".join(r.text for r in par.runs)
        new = full
        for k, v in variables.items():
            new = new.replace("{{" + k + "}}", str(v))
            new = new.replace("{{ " + k + " }}", str(v))
        if new != full and par.runs:
            par.runs[0].text = new
            for r in par.runs[1:]:
                r.text = ""

    for par in doc.paragraphs:
        repl_in_par(par)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    repl_in_par(par)
